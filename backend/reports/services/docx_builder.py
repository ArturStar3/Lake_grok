"""Сборка DOCX-отчётов через python-docx (те же collectors, что и для PDF)."""

from __future__ import annotations

import io
import re
from html.parser import HTMLParser

from django.utils import timezone
from django.utils.text import slugify
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from reports.services.collectors import collect_section_data

TABLE_SPECS = {
    'countries': (
        ('Страна', 'title'),
        ('Сокращение', 'title_short'),
        ('ISO', 'iso_code'),
    ),
    'targets': (
        ('Наименование', 'title'),
        ('Метка', 'label'),
        ('Страна', 'country'),
        ('Тип', 'type'),
        ('Широта', 'lat'),
        ('Долгота', 'lng'),
    ),
    'equipment': (
        ('Наименование', 'title'),
        ('Обозначение', 'designation'),
        ('Категория', 'category'),
        ('Страна происхождения', 'origin_country'),
        ('Описание', 'description'),
    ),
    'events': (
        ('Название', 'title'),
        ('Объект', 'object_name'),
        ('Тип', 'event_type'),
        ('Страна', 'country'),
        ('Дата', '_event_date'),
        ('Время', 'time_start'),
        ('Описание', 'description'),
    ),
    'situations': (
        ('Название', 'title'),
        ('Версия', 'version'),
        ('Дата', 'situation_date'),
        ('Время', 'situation_time'),
        ('Страны', 'countries'),
        ('Цвет', 'color'),
        ('Описание', 'description'),
    ),
    'zones': (
        ('Объект', 'target'),
        ('Страна', 'country'),
        ('Тип зоны', 'action_type'),
        ('Радиус, км', 'radius_km'),
        ('Геометрия', '_has_geometry'),
    ),
    'vulnerabilities': (
        ('Название', 'title'),
        ('Объект', 'target'),
        ('Страна', 'country'),
        ('Описание', 'description'),
        ('Широта', 'lat'),
        ('Долгота', 'lng'),
    ),
}


def _safe_filename(name: str) -> str:
    base = slugify(name, allow_unicode=True) or 'report'
    base = re.sub(r'[^\w\-]+', '_', base, flags=re.UNICODE).strip('_')
    return (base or 'report')[:80]


def _cell_value(row: dict, key: str):
    if key == '_event_date':
        start = row.get('date_start') or ''
        end = row.get('date_end') or ''
        if end and end != start:
            return f'{start} — {end}'
        return start
    if key == '_has_geometry':
        return 'Да' if row.get('has_geometry') else 'Нет'
    value = row.get(key, '')
    if value is None:
        return ''
    return str(value)


def _set_run_gray(run):
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    run.italic = True


def _set_paragraph_spacing(style, *, before_pt, after_pt, line_spacing=1.0, keep_with_next=False):
    """Explicit spacing so python-docx default template gaps do not accumulate."""
    pf = style.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if keep_with_next:
        pf.keep_with_next = True


def _tighten_styles(doc: Document):
    """
    Override default Word template spacing.

    Headings get at most ~2 lines of extra space; body text stays compact.
    """
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    _set_paragraph_spacing(normal, before_pt=0, after_pt=6, line_spacing=1.15)

    heading_specs = {
        'Title': (0, 10),
        'Heading 1': (14, 6),
        'Heading 2': (12, 4),
        'Heading 3': (10, 3),
        'Heading 4': (8, 2),
    }
    for style_name, (before_pt, after_pt) in heading_specs.items():
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        _set_paragraph_spacing(
            style,
            before_pt=before_pt,
            after_pt=after_pt,
            line_spacing=1.0,
            keep_with_next=True,
        )

    for style_name in ('List Bullet', 'List Number'):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        _set_paragraph_spacing(style, before_pt=0, after_pt=2, line_spacing=1.15)


def _write_denied(doc, text='Нет доступа к данным'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_gray(run)


def _write_empty(doc, text='Нет данных по выбранным фильтрам'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_gray(run)


def _write_rows_table(doc, columns, rows):
    """columns: list[(header, key)], rows: list[dict]."""
    if not rows:
        _write_empty(doc)
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for idx, (header, _) in enumerate(columns):
        header_cells[idx].text = header
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, (_, key) in enumerate(columns):
            cells[c_idx].text = _cell_value(row, key)


class _HtmlToDocxParser(HTMLParser):
    """Минимальный HTML → DOCX для контента формуляра/досье."""

    def __init__(self, doc: Document):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self._paragraph = None
        self._bold = 0
        self._italic = 0
        self._heading_level = None
        self._heading_parts: list[str] = []
        self._list_style = None
        self._in_li = False
        self._table_rows: list[list[str]] = []
        self._current_row: list[str] | None = None
        self._current_cell: list[str] | None = None
        self._skip_depth = 0

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in {'script', 'style'}:
            self._skip_depth += 1
            return
        if self._skip_depth:
            return
        if tag in {'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}:
            self._flush_paragraph()
            self._heading_level = int(tag[1])
            self._heading_parts = []
        elif tag == 'p':
            self._flush_paragraph()
            self._paragraph = self.doc.add_paragraph()
        elif tag == 'br':
            if self._paragraph is not None:
                self._paragraph.add_run().add_break()
            elif self._current_cell is not None:
                self._current_cell.append('\n')
        elif tag in {'strong', 'b'}:
            self._bold += 1
        elif tag in {'em', 'i'}:
            self._italic += 1
        elif tag == 'ul':
            self._flush_paragraph()
            self._list_style = 'List Bullet'
        elif tag == 'ol':
            self._flush_paragraph()
            self._list_style = 'List Number'
        elif tag == 'li':
            self._flush_paragraph()
            self._in_li = True
            style = self._list_style or 'List Bullet'
            self._paragraph = self.doc.add_paragraph(style=style)
        elif tag == 'table':
            self._flush_paragraph()
            self._table_rows = []
        elif tag == 'tr':
            self._current_row = []
        elif tag in {'td', 'th'}:
            self._current_cell = []
        elif tag == 'a':
            pass

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in {'script', 'style'}:
            self._skip_depth = max(0, self._skip_depth - 1)
            return
        if self._skip_depth:
            return
        if tag in {'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}:
            text = ''.join(self._heading_parts).strip()
            level = min(max(self._heading_level or 1, 1), 9)
            if text:
                self.doc.add_heading(text, level=level)
            self._heading_level = None
            self._heading_parts = []
        elif tag == 'p':
            self._flush_paragraph()
        elif tag in {'strong', 'b'}:
            self._bold = max(0, self._bold - 1)
        elif tag in {'em', 'i'}:
            self._italic = max(0, self._italic - 1)
        elif tag in {'ul', 'ol'}:
            self._list_style = None
        elif tag == 'li':
            self._flush_paragraph()
            self._in_li = False
        elif tag in {'td', 'th'}:
            if self._current_row is not None and self._current_cell is not None:
                self._current_row.append(''.join(self._current_cell).strip())
            self._current_cell = None
        elif tag == 'tr':
            if self._current_row is not None:
                self._table_rows.append(self._current_row)
            self._current_row = None
        elif tag == 'table':
            self._emit_table()
            self._table_rows = []

    def handle_data(self, data):
        if self._skip_depth or not data:
            return
        if self._heading_level is not None:
            self._heading_parts.append(data)
            return
        if self._current_cell is not None:
            self._current_cell.append(data)
            return
        if self._paragraph is None:
            # Loose text outside <p> — create a paragraph.
            self._paragraph = self.doc.add_paragraph()
        run = self._paragraph.add_run(data)
        if self._bold:
            run.bold = True
        if self._italic:
            run.italic = True

    def _flush_paragraph(self):
        self._paragraph = None

    def _emit_table(self):
        rows = [r for r in self._table_rows if r]
        if not rows:
            return
        cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Table Grid'
        for r_idx, row in enumerate(rows):
            for c_idx in range(cols):
                text = row[c_idx] if c_idx < len(row) else ''
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = text
                if r_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    def close(self):
        self._flush_paragraph()
        super().close()


def _html_to_docx(doc: Document, html_string: str):
    if not html_string or not str(html_string).strip():
        return
    parser = _HtmlToDocxParser(doc)
    parser.feed(str(html_string))
    parser.close()


def _write_dossier(doc: Document, dossier: dict | None):
    dossier = dossier or {}
    if dossier.get('access_denied'):
        _write_denied(doc, 'Нет доступа к данным')
        return
    standalone = dossier.get('standalone') or []
    groups = dossier.get('groups') or []
    if not standalone and not groups:
        _write_empty(doc, 'Нет заполненных разделов')
        return
    for item in standalone:
        title = item.get('section_title') or ''
        html = item.get('content_html') or ''
        if not html:
            continue
        if title:
            doc.add_heading(title, level=3)
        _html_to_docx(doc, html)
    for group in groups:
        parent = group.get('parent') or {}
        parent_title = parent.get('title') or ''
        if parent_title:
            doc.add_heading(parent_title, level=3)
        for child in group.get('children') or []:
            child_title = child.get('section_title') or ''
            html = child.get('content_html') or ''
            if child_title:
                doc.add_heading(child_title, level=4)
            if html:
                _html_to_docx(doc, html)


def _write_target_block(doc: Document, target: dict, *, page_break_before: bool = False):
    if page_break_before:
        doc.add_page_break()
    doc.add_heading(target.get('title') or 'Объект', level=2)
    meta_parts = []
    if target.get('type'):
        meta_parts.append(str(target['type']))
    if target.get('country'):
        meta_parts.append(str(target['country']))
    if target.get('label'):
        meta_parts.append(str(target['label']))
    lat = target.get('lat', '')
    lng = target.get('lng', '')
    meta_parts.append(f'{lat}, {lng}')
    doc.add_paragraph(' · '.join(meta_parts))

    doc.add_heading('Формуляр', level=3)
    formular = target.get('formular') or {}
    if formular.get('access_denied'):
        _write_denied(doc, target.get('formular_note') or 'Нет доступа к формуляру')
    else:
        _write_dossier(doc, formular)

    doc.add_heading('Зоны действия', level=3)
    zones = target.get('zones') or []
    if zones:
        _write_rows_table(
            doc,
            (
                ('Тип зоны', 'action_type'),
                ('Радиус, км', 'radius_km'),
                ('Геометрия', '_has_geometry'),
            ),
            zones,
        )
    else:
        _write_empty(doc, 'Зоны не заданы')

    doc.add_heading('Вооружение и техника', level=3)
    equipment = target.get('equipment') or []
    if equipment:
        _write_rows_table(
            doc,
            (
                ('Наименование', 'title'),
                ('Обозначение', 'designation'),
                ('Категория', 'category'),
                ('Кол-во', 'quantity'),
            ),
            equipment,
        )
    else:
        _write_empty(doc, 'Техника не указана')

    doc.add_heading('Уязвимости', level=3)
    vulns = target.get('vulnerabilities') or []
    if vulns:
        _write_rows_table(
            doc,
            (
                ('Название', 'title'),
                ('Описание', 'description'),
                ('Широта', 'lat'),
                ('Долгота', 'lng'),
            ),
            vulns,
        )
    else:
        _write_empty(doc, 'Уязвимости не указаны')


def _write_table_section(doc: Document, section_type: str, data: dict):
    if data.get('access_denied'):
        _write_denied(doc)
        return
    columns = TABLE_SPECS.get(section_type)
    if not columns:
        _write_empty(doc, 'Неизвестный тип раздела')
        return
    rows = data.get('rows') or []
    if not rows:
        _write_empty(doc)
        return
    _write_rows_table(doc, columns, rows)


def _write_country_full(doc: Document, data: dict):
    if data.get('access_denied'):
        _write_denied(doc)
        return
    countries = data.get('countries') or []
    if not countries:
        _write_empty(doc, 'Нет данных по выбранным странам')
        return
    for index, country in enumerate(countries):
        if index > 0:
            doc.add_page_break()
        doc.add_heading(country.get('title') or 'Страна', level=2)
        meta_parts = []
        if country.get('title_short'):
            meta_parts.append(str(country['title_short']))
        if country.get('iso_code'):
            meta_parts.append(str(country['iso_code']))
        if meta_parts:
            doc.add_paragraph(' · '.join(meta_parts))

        doc.add_heading('Досье страны', level=3)
        if country.get('dossier_access_denied'):
            _write_denied(doc, 'Нет доступа к досье страны')
        else:
            _write_dossier(doc, country.get('dossier'))

        doc.add_heading('Объекты', level=3)
        if country.get('targets_access_denied'):
            _write_denied(doc, 'Нет доступа к объектам')
        else:
            targets = country.get('targets') or []
            if not targets:
                _write_empty(doc, 'Объекты не найдены')
            else:
                for target in targets:
                    # Как в PDF: каждый объект начинается с новой страницы.
                    _write_target_block(doc, target, page_break_before=True)


def _write_objects_full(doc: Document, data: dict):
    if data.get('access_denied'):
        _write_denied(doc)
        return
    targets = data.get('targets') or []
    if not targets:
        _write_empty(doc, 'Нет данных по выбранным объектам')
        return
    for index, target in enumerate(targets):
        _write_target_block(doc, target, page_break_before=index > 0)


def _write_section(doc: Document, section, data: dict):
    title = getattr(section, 'title', None) or 'Раздел'
    doc.add_heading(title, level=1)
    section_type = section.section_type
    if section_type in TABLE_SPECS:
        _write_table_section(doc, section_type, data)
    elif section_type == 'country_full':
        _write_country_full(doc, data)
    elif section_type == 'objects_full':
        _write_objects_full(doc, data)
    else:
        _write_empty(doc, f'Тип раздела «{section_type}» не поддерживается в DOCX')


def build_docx_bytes(*, template_name, template_description, sections, user, section_overrides=None) -> bytes:
    overrides = section_overrides or {}
    ordered = sorted(list(sections), key=lambda s: (getattr(s, 'order', 0), getattr(s, 'id', 0) or 0))

    doc = Document()
    _tighten_styles(doc)

    doc.add_heading(template_name or 'Отчёт', level=0)
    if template_description:
        doc.add_paragraph(template_description)

    generated_at = timezone.localtime(timezone.now()).strftime('%d.%m.%Y %H:%M')
    user_name = ''
    if user is not None:
        full = (getattr(user, 'get_full_name', lambda: '')() or '').strip()
        user_name = full or getattr(user, 'username', '') or str(user)
    meta = f'Сформирован: {generated_at}'
    if user_name:
        meta = f'{meta} · {user_name}'
    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run(meta)
    _set_run_gray(meta_run)

    for index, section in enumerate(ordered):
        key = section.id if getattr(section, 'id', None) is not None else index
        filters = overrides.get(key, section.filters)
        data = collect_section_data(section.section_type, user, filters)
        if index > 0 and bool(getattr(section, 'page_break_before', True)):
            doc.add_page_break()
        _write_section(doc, section, data)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_docx_for_template(template, user, section_overrides=None) -> bytes:
    sections = list(template.sections.order_by('order', 'id'))
    return build_docx_bytes(
        template_name=template.name,
        template_description=template.description,
        sections=sections,
        user=user,
        section_overrides=section_overrides,
    )


def docx_filename(template_name: str) -> str:
    stamp = timezone.localtime(timezone.now()).strftime('%Y%m%d_%H%M')
    return f'{_safe_filename(template_name)}_{stamp}.docx'
